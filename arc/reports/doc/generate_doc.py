# TODO: arreglar y refactorizar
import datetime
import os
import re
import shutil
import time
import zipfile

import docx
from behave.runner import Context
from docx import table, Document
from docx.enum.dml import MSO_THEME_COLOR_INDEX
from docx.shared import RGBColor, Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from settings import settings
from arc.settings import settings as pytalos_settings
from arc.core.driver.driver_manager import DriverManager
from arc.core.driver.driver_setup import SetupDriver

BASE_PATH = os.path.abspath(os.path.join(os.path.abspath(__file__), os.pardir)) + os.sep
DOC_PATH = os.path.join(pytalos_settings.REPORTS_PATH, 'doc') + os.sep
TEMP_DOCX = BASE_PATH + "template.docx"
TEMP_ZIP = BASE_PATH + "template.zip"
TEMP_FOLDER = BASE_PATH + "template"


class GenerateDoc:
    document: Document
    scenario_name = ""
    cd = datetime.datetime.now()
    context: Context
    drive_type = ''
    resumen_table: table
    total_step_table: table
    step_info_doc = []
    step_table = None
    pytalos_config = None
    testcase_name = None

    def __init__(self, scenario):
        self.scenario = scenario
        default_wrapper = DriverManager.get_default_wrapper()
        config = default_wrapper.config.deepcopy()
        config_driver = SetupDriver(config)
        self.drive_type = config_driver.base_config.config.get('Driver', 'type')
        self.create_doc(scenario)
        self.step_info_doc = []
        self.add_note = ''

    def set_init_info(self, scenario):
        self.testcase_name = str(self.scenario_name) + ' - ' + str(self.drive_type).upper()
        self.document.add_heading(self.testcase_name, 1)
        self.generate_resumen_table(scenario)
        self.generate_total_step_table(scenario)
        self.document.add_page_break()

    def set_step_info_pre(self, step, extra_info_list, path=None, api_evidence_response=None, api_evidence_info=None,
                          api_evidence_headers=None, api_evidence_body=None, api_evidence_verify=None,
                          api_evidence_auto_extra_json=None, api_evidence_manual_extra_json=None):
        step_dict = {}
        result = self.format_status(step)
        cd = datetime.datetime.now()
        if result == "Passed":
            res_obt = "Operation with correct result"
        elif result == "Failed":
            res_obt = "Operation with incorrect result"
        else:
            res_obt = "Operation skipped"

        step_dict["Result"] = result
        step_dict["Date"] = cd.strftime('%d/%m/%Y %H:%M:%S').__str__()
        step_dict["Description"] = step.keyword.__str__() + " " + step.name.__str__()
        if step.text is not None:
            step_dict["Description"] = step_dict.get("Description") + "\n" + step.text.__str__()

        step_dict["Res-exp"] = step.name.__str__()
        step_dict["Res-obt"] = res_obt
        step_dict["Capture"] = path
        step_dict["Add_note"] = self.add_note
        step_dict["Extra_info_list"] = extra_info_list
        if str(self.drive_type).lower() in ['api', 'backend', 'service'] or "no_driver" in self.scenario.tags:
            step_dict["Api_Evidence_info"] = api_evidence_info
            step_dict["Api_Evidence_response"] = api_evidence_response
            step_dict["Api_Evidence_headers"] = api_evidence_headers
            step_dict["Api_Evidence_body"] = api_evidence_body
            step_dict["Api_Evidence_verify"] = api_evidence_verify
            step_dict["Api_Evidence_auto_extra_json"] = api_evidence_auto_extra_json
            step_dict["Api_Evidence_manual_extra_json"] = api_evidence_manual_extra_json

        if step.table is not None:
            step_dict["Extra_info"] = [True, step]
        else:
            step_dict["Extra_info"] = [False, step]

        self.step_info_doc.append(step_dict)

    def end(self, context):
        self.context = context
        self.genera_doc_body()
        aux = re.sub('[\'#%<>@\" ]', "", self.scenario_name)
        doc_path = re.sub('[.]', "-", aux)
        path = str(DOC_PATH) + str(self.drive_type).upper() + "-" + doc_path.__str__() + ".docx"
        self.document.save(path)

        return path

    def set_templeate_header(self):
        # remove old zip file or folder template
        if os.path.exists(TEMP_ZIP):
            os.remove(TEMP_ZIP)
        if os.path.exists(TEMP_FOLDER):
            shutil.rmtree(TEMP_FOLDER)

        dir_path = os.path.abspath(os.path.join(os.path.abspath(__file__), os.pardir)) + os.sep
        shutil.copy(dir_path + 'template.docx', dir_path + 'template_temp.docx')

        # reformat template.docx's extension
        os.rename(TEMP_DOCX, TEMP_ZIP)
        # unzip file zip to specific folder
        with zipfile.ZipFile(TEMP_ZIP, 'r') as z:
            z.extractall(TEMP_FOLDER)

        # change header xml file
        header_xml = os.path.join(TEMP_FOLDER, "word", "header2.xml")
        xmlstring = open(header_xml, 'r', encoding='utf-8').read()
        xmlstring = xmlstring.replace("tttt", "Test Case:\n " + str(self.scenario_name))
        currentdate = str(self.cd.day) + "/" + str(self.cd.month) + "/" + str(self.cd.year)
        xmlstring = xmlstring.replace("fff", currentdate)

        with open(header_xml, "wb") as f:
            f.write(xmlstring.encode("UTF-8"))

        # zip temp folder to zip file
        os.remove(TEMP_ZIP)
        shutil.make_archive(TEMP_ZIP.replace(".zip", ""), 'zip', TEMP_FOLDER)

        # rename zip file to docx
        os.rename(TEMP_ZIP, TEMP_DOCX)
        shutil.rmtree(TEMP_FOLDER)

    @staticmethod
    def remove_tem_files():
        dir_path = os.path.abspath(os.path.join(os.path.abspath(__file__), os.pardir)) + os.sep
        os.remove(dir_path + 'template.docx')
        os.rename(dir_path + 'template_temp.docx', dir_path + 'template.docx')

    def create_doc(self, scenario):
        file_path = BASE_PATH + 'template.docx'
        self.scenario_name = str(scenario.name)
        self.set_templeate_header()
        template = open(file_path, 'rb')
        self.document = Document(template)
        template.close()
        self.remove_tem_files()

    def format_subfile_name(self, scenario):
        sub = str(scenario.feature.filename).split("/")
        subfile_name = ""
        for name in sub:
            if name != 'features':
                subfile_name = subfile_name + name.capitalize() + "-"
        subfile_name = subfile_name.replace(".feature", "")
        return subfile_name

    def generate_resumen_table(self, scenario):
        fd = datetime.datetime.now()
        table_r = self.document.add_table(rows=11, cols=2)
        self.resumen_table = table_r
        table_r.style = 'Tablanormal1'
        cella = table_r.cell(0, 0)
        cellb = table_r.cell(0, 1)
        cell_merge = cella.merge(cellb)
        cell = table_r.cell(0, 0)
        self.add_bookmark(cell.paragraphs[0], 'Test Case Execution Summary', "ResumenTable")
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.set_result_format(table_r)

        table_r.cell(1, 0).text = 'Result'
        table_r.cell(2, 0).text = 'Number of steps'
        table_r.cell(3, 0).text = 'Correct steps'
        table_r.cell(4, 0).text = 'Incorrect steps'
        table_r.cell(5, 0).text = 'No Run steps'
        table_r.cell(6, 0).text = 'Start'
        table_r.cell(7, 0).text = 'End'
        table_r.cell(8, 0).text = 'Duration'
        table_r.cell(9, 0).text = 'Execution Type'
        table_r.cell(10, 0).text = 'Environment'

        if str(scenario.status) == 'Status.passed':
            result = "Passed"
        elif str(scenario.status) == "Status.failed":
            result = "Failed"
        else:
            result = "No Run"

        steps = []
        steps_passed = []
        steps_failed = []
        steps_skip = []
        for step in scenario.all_steps:
            steps.append(step)
            if step.status.__str__() == 'Status.passed':
                steps_passed.append(step)
            if step.status.__str__() == 'Status.failed':
                steps_failed.append(step)
            else:
                steps_skip.append(step)

        no_run = steps.__len__() - steps_passed.__len__() - steps_failed.__len__()
        cell = table_r.cell(1, 1)
        cell.text = result
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(2, 1)
        cell.text = steps.__len__().__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(3, 1)
        cell.text = steps_passed.__len__().__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(4, 1)
        cell.text = steps_failed.__len__().__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(5, 1)
        cell.text = no_run.__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(6, 1)
        cell.text = self.cd.strftime('%d/%m/%Y %H:%M:%S').__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(7, 1)
        cell.text = fd.strftime('%d/%m/%Y %H:%M:%S').__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(8, 1)
        test_duration = time.strftime('%H:%M:%S', time.gmtime(scenario.duration))
        cell.text = test_duration.__str__()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(9, 1)
        cell.text = str(self.drive_type).capitalize()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        cell = table_r.cell(10, 1)
        cell.text = str(settings.PYTALOS_PROFILES['environment']).upper()
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        self.document.add_paragraph('')
        self.set_result_format(self.resumen_table)

    def generate_total_step_table(self, scenario):
        steps = []
        steps_passed = []
        steps_failed = []
        steps_skip = []
        for step in scenario.all_steps:
            steps.append(step)
            if step.status.__str__() == 'Status.passed':
                steps_passed.append(step)
            if step.status.__str__() == 'Status.failed':
                steps_failed.append(step)
            else:
                steps_skip.append(step)

        rows = steps.__len__()
        table_r = self.document.add_table(rows=rows + 1, cols=4)
        self.total_step_table = table_r
        table_r.style = 'Tablanormal1'

        cell = table_r.cell(0, 0)
        cell.text = 'Steps'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cell = table_r.cell(0, 1)
        cell.text = 'Description'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cell = table_r.cell(0, 2)
        cell.text = 'Result'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cell = table_r.cell(0, 3)
        cell.text = 'Duration'
        cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

        cont_row = 1
        for row in steps:
            cell = table_r.cell(cont_row, 0)
            step_text = "Step " + str(cont_row)
            self.add_link(cell.paragraphs[0], "Step" + str(cont_row) + "bm", step_text)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell = table_r.cell(cont_row, 1)
            cell.text = row.name.__str__()
            cell = table_r.cell(cont_row, 2)
            cell.text = self.format_status(row)
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell = table_r.cell(cont_row, 3)
            test_duration = time.strftime('%H:%M:%S', time.gmtime(row.duration))
            cell.text = test_duration.__str__()
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cont_row += 1

        self.set_column_width(table_r)
        self.set_result_format(table_r)

    @staticmethod
    def set_column_width(table_width):
        for cell in table_width.columns[0].cells:
            cell.width = Inches(1.0)
        for cell in table_width.columns[1].cells:
            cell.width = Inches(4.0)
        for cell in table_width.columns[2].cells:
            cell.width = Inches(1.0)
        for cell in table_width.columns[3].cells:
            cell.width = Inches(1.0)

    @staticmethod
    def set_result_format(table_in):
        for row in table_in.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text == 'Failed':
                            font = run.font
                            font.color.rgb = RGBColor(252, 3, 3)
                            font.bold = True
                        if run.text == 'Passed':
                            font = run.font
                            font.color.rgb = RGBColor(3, 156, 51)
                            font.bold = True
                        if run.text == 'No Run':
                            font = run.font
                            font.color.rgb = RGBColor(0, 108, 255)
                            font.bold = True
                        if run.text == 'Error':
                            font = run.font
                            font.color.rgb = RGBColor(191, 13, 13)
                            font.bold = True
                        if run.text == 'Description:' or run.text == 'Expected result:' or run.text == 'Obtained result:' or run.text == 'Data Table':
                            font = run.font
                            font.color.rgb = RGBColor(0, 0, 0)
                            font.bold = True
                        if "Nota: " in run.text:
                            font = run.font
                            font.color.rgb = RGBColor(0, 0, 0)
                            font.bold = True
                        else:
                            font = run.font
                            font.bold = False

    @staticmethod
    def format_status(step):
        if step.status.__str__() == 'Status.passed':
            return 'Passed'
        if step.status.__str__() == 'Status.failed':
            return 'Failed'
        else:
            return 'No Run'

    def genera_doc_body(self):
        cont = 1
        for step in self.step_info_doc:
            table_r = self.document.add_table(rows=1, cols=4)
            table_r.style = 'Tablanormal1'
            cell = table_r.cell(0, 0)
            step_text = "Step " + cont.__str__()
            self.add_bookmark(cell.paragraphs[0], step_text, "Step" + cont.__str__() + "bm")
            cont += 1
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell = table_r.cell(0, 1)
            cell.text = step.get("Result")
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell = table_r.cell(0, 2)
            cell.text = step.get("Date")
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            cell = table_r.cell(0, 3)
            self.add_link(cell.paragraphs[0], "ResumenTable", "Go to Summary Table")
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

            for cell in table_r.columns[0].cells:
                cell.width = Inches(1.0)
            for cell in table_r.columns[1].cells:
                cell.width = Inches(1.0)
            for cell in table_r.columns[2].cells:
                cell.width = Inches(3.0)
            for cell in table_r.columns[3].cells:
                cell.width = Inches(3.0)

            self.set_result_format(table_r)

            self.document.add_heading('Description:', 2)
            p = self.document.add_paragraph(step.get('Description'))

            if step.get('Extra_info')[0]:
                self.set_data_table(step.get('Extra_info')[1])

            self.document.add_heading('Expected result:', 2)
            p = self.document.add_paragraph(step.get('Res-exp'))
            self.document.add_heading('Obtained result:', 2)
            p = self.document.add_paragraph(step.get('Res-obt'))
            self.document.add_paragraph('')

            if step.get('Res-obt') == 'Operation with incorrect result':
                step_error = step.get('Extra_info')[1]
                if step_error.status == "failed":
                    txt_ex = step_error.exception
                else:
                    txt_ex = None

                if txt_ex is not None or txt_ex != '':
                    table_r = self.document.add_table(rows=2, cols=1)
                    table_r.style = 'Tablanormal1'
                    cell = table_r.cell(0, 1)
                    cell.text = str(txt_ex)
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    cell = table_r.cell(0, 0)
                    cell.text = "Error"
                    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
                    self.set_result_format(table_r)

            extra_info_list = step.get("Extra_info_list")
            text_info = step.get("Add_note")
            try:
                if text_info != "":
                    self.document.add_heading('Nota: ', 5)
                    p = self.document.add_paragraph(str(text_info))
                for info in extra_info_list:
                    if ".png" in info:
                        if str(self.drive_type).lower() == "android" or str(self.drive_type).lower() == "ios":
                            self.document.add_picture(step.get('Capture'), width=Inches(3.0))
                            last_paragraph = self.document.paragraphs[-1]
                            last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                        else:
                            self.document.add_picture(step.get('Capture'), width=Inches(6.5))
                    else:
                        self.document.add_paragraph(info)
                if str(self.drive_type).lower() == "android" or str(self.drive_type).lower() == "ios":
                    self.document.add_picture(step.get('Capture'), width=Inches(3.0))
                    last_paragraph = self.document.paragraphs[-1]
                    last_paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                elif str(self.drive_type).lower() in ['api', 'backend', 'service']:
                    pass
                else:
                    self.document.add_picture(step.get('Capture'), width=Inches(6.5))
                self.document.add_paragraph('')

            except FileNotFoundError:
                self.document.add_paragraph('')

            if str(self.drive_type).lower() in ['api', 'backend', 'service'] or "no_driver" in self.scenario.tags:
                self.add_api_evidence_body(step)

            self.document.add_paragraph('')
            self.document.add_page_break()

    def add_api_evidence_body(self, step):
        # Request Info
        self.add_api_evidence_request_info(step)

        # Request Headers
        self.add_api_evidence_request_headers(step)

        # Request Body
        self.add_api_evidence_request_body(step)

        # JSON Response
        self.add_api_evidence_json_response(step)

        # Verify Evidence
        self.add_api_evidence_verify(step)

        # Set Extra Json Auto Evidence
        self.add_api_evidence_auto_extra_json(step)

        # Set Extra Json Manual Evidence
        self.add_api_evidence_manual_extra_json(step)

    def add_api_evidence_verify(self, step):
        if step.get("Api_Evidence_verify"):
            for dictionary in step.get("Api_Evidence_verify"):
                len_dict = len(dictionary)
                info_dict = dictionary
                row_cont = len_dict
                if info_dict.get("Error Message") is None:
                    row_cont -= 1

                table_r = self.document.add_table(rows=row_cont, cols=2)
                table_r.style = 'Tablanormal1'
                cella = table_r.cell(0, 0)
                cellb = table_r.cell(0, 1)
                cell_merge = cella.merge(cellb)
                cell_merge.text = info_dict["Verify_Name"]
                cell = table_r.cell(0, 0)

                cont = 1
                for key in info_dict.keys():
                    if key != "Verify_Name" and key != "Error Message":
                        if info_dict.get(key) is not None:
                            cell = table_r.cell(cont, 0)
                            cell.text = str(key)
                            cell = table_r.cell(cont, 1)
                            cell.text = str(info_dict.get(key))
                            cont += 1
                    elif key == "Error Message":
                        if info_dict.get(key) is not None:
                            cell = table_r.cell(cont, 0)
                            cell.text = str(key)
                            cell = table_r.cell(cont, 1)
                            cell.text = str(info_dict.get(key))
                            cont += 1

                self.set_api_evidence_key_format(table_r)
                for row in table_r.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text == str(info_dict["Verify_Name"]):
                                    font = run.font
                                    font.color.rgb = RGBColor(252, 3, 3)
                                    font.bold = True
                                else:
                                    font = run.font
                                    font.bold = False
                self.set_result_format(table_r)
                self.document.add_paragraph('')

    def add_api_evidence_request_info(self, step):
        if step.get("Api_Evidence_info"):
            table_r = self.document.add_table(rows=5, cols=2)
            table_r.style = 'Tablanormal1'
            cella = table_r.cell(0, 0)
            cellb = table_r.cell(0, 1)
            cell_merge = cella.merge(cellb)
            cell_merge.text = 'Request Info'
            cell = table_r.cell(0, 0)

            info_dict = step.get("Api_Evidence_info")
            cell = table_r.cell(1, 0)
            cell.text = "URL"
            cell = table_r.cell(2, 0)
            cell.text = "Method"
            cell = table_r.cell(3, 0)
            cell.text = "Status Code"
            cell = table_r.cell(4, 0)
            cell.text = "Remote Address"

            cell = table_r.cell(1, 1)
            cell.text = info_dict.get("URL")
            cell = table_r.cell(2, 1)
            cell.text = str(info_dict.get("Method")).upper()
            cell = table_r.cell(3, 1)
            cell.text = str(info_dict.get("Status Code")) + " " + info_dict.get("Reason")
            cell = table_r.cell(4, 1)
            cell.text = str(info_dict.get("Remote Address"))
            for cell in table_r.columns[0].cells:
                cell.width = Inches(1.5)
            for cell in table_r.columns[1].cells:
                cell.width = Inches(5.5)

            self.set_api_evidence_key_format(table_r)
            self.set_api_evidence_title_colour(table_r)
            self.document.add_paragraph('')

    def add_api_evidence_request_headers(self, step):
        if step.get("Api_Evidence_headers"):
            headers_datas = step.get("Api_Evidence_headers")
            count_data_index = len(headers_datas)
            table_r = self.document.add_table(rows=count_data_index + 1, cols=2)
            table_r.style = 'Tablanormal1'
            cella = table_r.cell(0, 0)
            cellb = table_r.cell(0, 1)
            cell_merge = cella.merge(cellb)
            cell_merge.text = 'Request Headers'
            cell = table_r.cell(0, 0)

            cont = 1
            headers_dict = headers_datas
            for key in headers_dict.keys():
                cell = table_r.cell(cont, 0)
                cell.text = str(key).capitalize()
                cell = table_r.cell(cont, 1)
                cell.text = headers_dict[key]
                cont += 1
            self.set_api_evidence_key_format(table_r)
            self.set_api_evidence_title_colour(table_r)
            self.document.add_paragraph('')

    def add_api_evidence_json_response(self, step):
        if step.get("Api_Evidence_response") and str(step.get("Api_Evidence_response")) != "null":
            api_evidence = step.get("Api_Evidence_response")
            table_r = self.document.add_table(rows=2, cols=1)
            table_r.style = 'Tablanormal1'
            cell = table_r.cell(0, 0)
            cell.text = "JSON Response"

            cell = table_r.cell(1, 0)
            cell.text = api_evidence
            for row in table_r.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            if run.text == "JSON Response":
                                font = run.font
                                font.bold = True
                            else:
                                font = run.font
                                font.bold = False
            self.set_api_evidence_title_colour(table_r)
            self.document.add_paragraph('')

    def add_api_evidence_auto_extra_json(self, step):
        for inte in step.get("Api_Evidence_auto_extra_json"):
            if inte and str(inte) != "null":
                api_evidence = inte
                table_r = self.document.add_table(rows=2, cols=1)
                table_r.style = 'Tablanormal1'
                cell = table_r.cell(0, 0)
                cell.text = "Response Headers"

                cell = table_r.cell(1, 0)
                cell.text = api_evidence
                for row in table_r.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text == "Response Headers":
                                    font = run.font
                                    font.bold = False
                                else:
                                    font = run.font
                                    font.bold = False
                self.set_api_evidence_title_colour(table_r)
                self.document.add_paragraph('')

    def add_api_evidence_manual_extra_json(self, step):
        for inte in step.get("Api_Evidence_manual_extra_json"):
            if inte and str(inte) != "null":
                title = inte.get("TITLE_EVIDENCE")
                try:
                    del inte['TITLE_EVIDENCE']
                except Exception:
                    pass
                table_r = self.document.add_table(rows=2, cols=1)
                table_r.style = 'Tablanormal1'
                cell = table_r.cell(0, 0)
                cell.text = title
                evidence_text = self.context.func.format_evidence_json(inte)
                cell = table_r.cell(1, 0)

                cell.text = evidence_text
                for row in table_r.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text == title:
                                    font = run.font
                                    font.color.rgb = RGBColor(252, 3, 3)
                                    font.bold = False
                                else:
                                    font = run.font
                                    font.bold = False
                self.set_api_evidence_title_colour(table_r)
                self.document.add_paragraph('')

    def add_api_evidence_request_body(self, step):
        for inte in step.get("Api_Evidence_body"):
            if inte and str(inte) != "null":
                try:
                    del inte['TITLE_EVIDENCE']
                except Exception:
                    pass
                table_r = self.document.add_table(rows=2, cols=1)
                table_r.style = 'Tablanormal1'
                cell = table_r.cell(0, 0)
                cell.text = "Request Body"
                evidence_text = self.context.func.format_evidence_json(inte)
                cell = table_r.cell(1, 0)

                cell.text = evidence_text
                for row in table_r.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            for run in paragraph.runs:
                                if run.text == "Request Body":
                                    font = run.font
                                    font.bold = False
                                else:
                                    font = run.font
                                    font.bold = False
                self.set_api_evidence_title_colour(table_r)
                self.document.add_paragraph('')

    @staticmethod
    def set_api_evidence_title_colour(table_in):
        for row in table_in.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text == 'Request Info' or run.text == 'Request Headers' or run.text == 'JSON Response' or run.text == 'Request Body' or run.text == 'Response Headers':
                            font = run.font
                            font.color.rgb = RGBColor(252, 3, 3)
                            font.bold = False

    @staticmethod
    def set_api_evidence_key_format(table_in):
        for row in table_in.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        font = run.font
                        font.bold = False

    def add_screenshot(self, path):
        try:
            self.document.add_picture(path, width=Inches(6.5))
        except FileNotFoundError:
            self.document.add_paragraph('')

    def set_data_table(self, step):
        cols = len(step.table.headings)
        rows = len(step.table.rows)
        table_r = self.document.add_table(rows=rows + 2, cols=cols)
        table_r.style = 'Tablanormal1'

        cella = table_r.cell(0, 0)
        cellb = table_r.cell(0, cols - 1)
        cell_merge = cella.merge(cellb)
        cell_merge.text = 'Data Table'
        cell = table_r.cell(0, 0)

        cont = 0
        for head in step.table.headings:
            cell = table_r.cell(1, cont)
            cell.text = head.__str__()
            cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            cont += 1

        cont_rows = 2
        for key in step.table:
            cont_cols = 0
            for cells in key:
                cell = table_r.cell(cont_rows, cont_cols)
                cell.text = cells.__str__()
                cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                cont_cols += 1
            cont_rows += 1

        for row in table_r.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        if run.text == 'Data Table':
                            font = run.font
                            font.bold = True
                        else:
                            font = run.font
                            font.bold = False

    def add_bookmark(self, paragraph, bookmark_text, bookmark_name):
        run = paragraph.add_run()
        tag = run._r
        start = docx.oxml.shared.OxmlElement('w:bookmarkStart')
        start.set(docx.oxml.ns.qn('w:id'), '0')
        start.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(start)

        text = docx.oxml.OxmlElement('w:r')
        text.text = bookmark_text
        tag.append(text)

        end = docx.oxml.shared.OxmlElement('w:bookmarkEnd')
        end.set(docx.oxml.ns.qn('w:id'), '0')
        end.set(docx.oxml.ns.qn('w:name'), bookmark_name)
        tag.append(end)

    def add_link(self, paragraph, link_to, text, tool_tip=None):
        # create hyperlink node
        hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')

        # set attribute for link to bookmark
        hyperlink.set(docx.oxml.shared.qn('w:anchor'), link_to, )

        if tool_tip is not None:
            # set attribute for link to bookmark
            hyperlink.set(docx.oxml.shared.qn('w:tooltip'), tool_tip, )

        new_run = docx.oxml.shared.OxmlElement('w:r')
        rPr = docx.oxml.shared.OxmlElement('w:rPr')
        new_run.append(rPr)
        new_run.text = text
        hyperlink.append(new_run)
        r = paragraph.add_run()
        r._r.append(hyperlink)
        r.font.name = "Calibri"
        r.font.color.theme_color = MSO_THEME_COLOR_INDEX.HYPERLINK
        r.font.underline = True